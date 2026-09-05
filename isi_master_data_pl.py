"""isi_master_data_pl.py — Auto-isi sheet '@ Master Data' paket PL via COM.

Trigger macro VBA `ModDraftPaketPL.IsiDataPLByKode` (reuse 100% logika VBA:
lookup KPA/Dinas, compose nomor dokpil/undangan/BA reviu, kode unik, personil).

Dipanggil dari app.py saat buat folder PL (single/bulk), setelah HPS + refresh
template. Menggantikan tombol manual "Muat Paket PL" + "Isi Data PL" di Excel.
"""

import os
from datetime import date, datetime, timedelta

from person_name_utils import normalize_equipment_quantity


K3_CERT_FALLBACK = "SKK Petugas K3 Konstruksi/Keselamatan Konstruksi"
PAYMENT_METHOD_MC = (
    "Monthly Certificate/MC dimana pembayaran prestasi pekerjaan dilakukan "
    "secara bulanan berdasarkan prestasi pekerjaan yang telah dicapai, yang "
    "dibuktikan dengan hasil pengukuran bersama dan Berita Acara Pemeriksaan "
    "Hasil Pekerjaan"
)
PAYMENT_METHOD_TERMIN = "Termin"


def _coerce_create_folder_date(value=None) -> date:
    """Normalisasi tanggal yang ditangkap sekali saat tombol create-folder diklik."""
    if value is None:
        return datetime.now().date()
    if isinstance(value, datetime):
        return value.date()
    if isinstance(value, date):
        return value
    raise TypeError(f"tanggal create-folder tidak valid: {value!r}")


def set_create_folder_dates(workbook, created_on=None, progress_cb=None) -> date:
    """Isi tanggal BA Reviu dan Dokpil pada boundary workbook PL.

    ``H8:H10`` adalah sumber hari/bulan/tahun BA Reviu, sedangkan ``C21``
    adalah tanggal Dokpil. C21 sengaja ditulis sebagai nilai tanggal COM,
    bukan string, supaya formula dan mail merge tetap menerima tanggal Excel.
    """
    created = _coerce_create_folder_date(created_on)
    ws = workbook.Worksheets("@ Master Data")

    for address, value in (
        ("H8", created.day),
        ("H9", created.month),
        ("H10", created.year),
    ):
        cell = ws.Range(address)
        cell.NumberFormat = "General"
        cell.Value = value

    dokpil_cell = ws.Range("C21")
    dokpil_cell.NumberFormat = "[$-id-ID]dd mmmm yyyy"
    # Value2 memakai serial Excel langsung. Mengirim Python datetime melalui
    # COM dapat bergeser satu hari akibat konversi timezone regional Windows.
    excel_epoch = date(1899, 12, 30)
    dokpil_cell.Value2 = (created - excel_epoch).days

    # Fail closed: jangan menyimpan workbook jika Excel tidak menerima nilai
    # target secara utuh. Tidak memanggil Calculate agar cache UDF tetap aman.
    for address, expected in (
        ("H8", created.day),
        ("H9", created.month),
        ("H10", created.year),
    ):
        actual = ws.Range(address).Value
        if int(float(actual)) != expected:
            raise RuntimeError(
                f"verifikasi {address} gagal: expected={expected!r}, actual={actual!r}"
            )

    actual_date = dokpil_cell.Value2
    if isinstance(actual_date, datetime):
        actual_date = actual_date.date()
    elif isinstance(actual_date, date):
        pass
    elif all(hasattr(actual_date, part) for part in ("year", "month", "day")):
        actual_date = date(
            int(actual_date.year), int(actual_date.month), int(actual_date.day)
        )
    elif isinstance(actual_date, (int, float)):
        actual_date = (datetime(1899, 12, 30) + timedelta(days=float(actual_date))).date()
    else:
        actual_date = None
    if actual_date != created:
        raise RuntimeError(
            f"verifikasi C21 gagal: expected={created!r}, actual={actual_date!r}"
        )

    if progress_cb:
        progress_cb(
            "Tanggal create-folder: "
            f"H8:H10={created.day:02d}/{created.month:02d}/{created.year}; "
            f"C21={created.isoformat()}"
        )
    return created


def _is_plpk_workbook(excel_path: str, family: str = "") -> bool:
    """True hanya untuk workbook PLPK; enrichment tidak boleh masuk PLJKK."""
    normalized_family = str(family or "").strip().upper()
    if normalized_family in {"JKK", "PLJKK"}:
        return False
    if normalized_family in {"PK", "PLPK"}:
        return True
    workbook_name = os.path.basename(str(excel_path or "")).upper()
    parent_name = os.path.basename(os.path.dirname(os.path.abspath(str(excel_path or "")))).upper()
    return "BAPLPK" in workbook_name or "PLPK" in parent_name


def infer_cara_pembayaran_plpk(folder: str, package_title: str = "") -> str:
    """Tentukan default cara pembayaran PLPK dari konteks paket lokal.

    Infrastruktur PUPR (atau judul pekerjaan jalan, drainase, box culvert,
    jembatan, dan pekerjaan sejenis) memakai Monthly Certificate/MC. Paket
    lain memakai Termin. Resolver hanya membaca nama folder/file, sehingga
    tidak menambah fetch jaringan atau bergantung pada cache Supabase.
    """
    import re

    context_parts = [str(package_title or ""), os.path.basename(os.path.normpath(folder or ""))]
    if folder and os.path.isdir(folder):
        try:
            for root, dirs, files in os.walk(folder):
                context_parts.extend(dirs)
                context_parts.extend(files)
        except OSError:
            pass
    context = " ".join(context_parts).casefold()

    is_pupr = bool(re.search(
        r"\b(?:dpupr|pupr|bina\s+marga|pekerjaan\s+umum|dinas\s+pu)\b",
        context,
    ))
    is_infrastructure = bool(re.search(
        r"\b(?:jalan|jembatan|drainase|box\s*culvert|rabat\s+beton|"
        r"aspal|latasir|makadam)\b",
        context,
    ))
    return PAYMENT_METHOD_MC if is_pupr or is_infrastructure else PAYMENT_METHOD_TERMIN


def _normalize_contract_type(value: str) -> str:
    """PLPK hanya memakai Lumsum atau Harga Satuan."""
    import re
    text = re.sub(r"\s+", " ", str(value or "")).strip()
    lowered = text.lower()
    if "gabungan" in lowered or (
        "lumsum" in lowered and "harga satuan" in lowered
    ):
        return "Harga Satuan"
    if "harga satuan" in lowered:
        return "Harga Satuan"
    if lowered == "lumsum":
        return "Lumsum"
    return text


def _normalize_k3_certificate(value: str, jabatan: str = "") -> str:
    """Selaraskan sertifikat dengan bunyi jabatan K3 pada sumber paket."""
    import re
    text = re.sub(r"\s+", " ", str(value or "")).strip()
    role = re.sub(r"\s+", " ", str(jabatan or "")).strip()
    if role:
        if re.search(r"\bpetugas\s+k3\s+konstruksi\b", role, re.IGNORECASE):
            if re.search(r"\bSKK\s+Petugas\s+K3\s+Konstruksi\b", text, re.IGNORECASE) and not re.search(
                r"keselamatan\s+konstruksi", text, re.IGNORECASE
            ):
                return text
            return "SKK Petugas K3 Konstruksi"
        if re.search(r"\bpetugas\s+k3\b", role, re.IGNORECASE):
            return K3_CERT_FALLBACK
    explicit = re.search(r"\bSKK\b", text, re.IGNORECASE) and re.search(
        r"petugas\s+k3|keselamatan\s+konstruksi", text, re.IGNORECASE
    )
    return text if explicit else K3_CERT_FALLBACK


def _clean_docx_cell(value: str) -> str:
    """Normalisasi teks sel DOCX tanpa mengubah makna isinya."""
    import re

    return (
        re.sub(r"\s+", " ", str(value or ""))
        .replace("\ufffd", " ")
        .replace("\u2019", "'")
        .strip()
    )


def _equipment_header(table):
    """Cari header tabel alat dan indeks kolomnya berdasarkan label."""
    import re

    for header_row, row in enumerate(table.rows[:5]):
        headers = [_clean_docx_cell(cell.text).casefold() for cell in row.cells]
        name_col = next(
            (
                i
                for i, value in enumerate(headers)
                if ("peralatan" in value or re.search(r"\balat\b", value))
                and ("nama" in value or "jenis" in value or "perlengkapan" in value)
            ),
            None,
        )
        capacity_col = next(
            (i for i, value in enumerate(headers) if re.search(r"\bkapasitas\b", value)),
            None,
        )
        quantity_col = next(
            (i for i, value in enumerate(headers) if re.search(r"\bjumlah\b", value)),
            None,
        )
        if name_col is not None and capacity_col is not None and quantity_col is not None:
            return header_row, name_col, capacity_col, quantity_col
    return None


def _parse_equipment_docx(path: str) -> list[dict]:
    """Parse tabel peralatan DOCX; deteksi berdasarkan header, bukan nama file."""
    import re
    from docx import Document

    result = []
    seen = set()
    doc = Document(path)
    for table in doc.tables:
        header = _equipment_header(table)
        if header is None:
            continue
        header_row, name_col, capacity_col, quantity_col = header
        for row in table.rows[header_row + 1 :]:
            raw_values = [cell.text or "" for cell in row.cells]
            values = [_clean_docx_cell(value) for value in raw_values]
            if not values or not re.match(r"^\s*\d+(?:\s*[.)-]|\s*$)", values[0]):
                continue
            if name_col >= len(values):
                continue
            names = [_clean_docx_cell(part) for part in raw_values[name_col].splitlines() if part.strip()]
            if not names:
                continue
            capacities = (
                [_clean_docx_cell(part) for part in raw_values[capacity_col].splitlines() if part.strip()]
                if capacity_col < len(values)
                else []
            )
            quantities = (
                [_clean_docx_cell(part) for part in raw_values[quantity_col].splitlines() if part.strip()]
                if quantity_col < len(values)
                else []
            )
            for index, name in enumerate(names):
                if name.casefold() in {
                    "nama peralatan",
                    "jenis peralatan",
                    "jenis peralatan/perlengkapan",
                    "-",
                }:
                    continue
                capacity = capacities[index] if index < len(capacities) else ""
                quantity = quantities[index] if index < len(quantities) else ""
                identity = (name.casefold(), capacity.casefold(), quantity.casefold())
                if identity in seen:
                    continue
                seen.add(identity)
                result.append({
                    "nama": name,
                    "kapasitas": capacity,
                    "jumlah": normalize_equipment_quantity(quantity),
                })
    return result


def _parse_local_enrichment(folder: str) -> dict:
    """Ambil field yang memang bersumber dari artefak lokal paket.

    Field ini tidak bergantung schema Supabase: alat, uraian HPS, dan ringkasan
    RK3 langsung ditulis ke workbook setelah macro master-data selesai.
    """
    import re
    out = {"personil": [], "alat": [], "uraian": [], "uraian_rk3": [], "risiko": "", "risiko_tertinggi": "", "contract_type": "", "provider": {}, "cara_pembayaran": infer_cara_pembayaran_plpk(folder),}
    try:
        import parse_kak_pl as pk
        person = pk.cari_daftar_personil_di_folder(folder)
        if person:
            out["personil"] = pk.parse_personil_daftar(person)
        nd = pk.cari_nd_di_folder(folder)
        if nd:
            out["provider"] = pk.parse_nd_penyedia(nd)
        draft = pk.cari_draft_pl_di_folder(folder)
        if draft:
            out["contract_type"] = pk.parse_jenis_kontrak_dari_draft_pl(draft)
    except Exception:
        pass
    try:
        from docx import Document
        for root, _, files in os.walk(folder):
            for name in files:
                low = name.lower()
                path = os.path.join(root, name)
                if low.endswith(".docx"):
                    # Jangan bergantung pada nama file; dokumen KAK dapat
                    # memakai nama umum untuk tabel alat yang valid.
                    try:
                        for item in _parse_equipment_docx(path):
                            identity = (
                                item["nama"].casefold(),
                                item["kapasitas"].casefold(),
                                item["jumlah"].casefold(),
                            )
                            if not any(
                                (
                                    existing["nama"].casefold(),
                                    existing["kapasitas"].casefold(),
                                    existing["jumlah"].casefold(),
                                )
                                == identity
                                for existing in out["alat"]
                            ):
                                out["alat"].append(item)
                    except Exception:
                        # Dokumen non-KAK atau DOCX rusak tidak boleh
                        # menghentikan pemindaian dokumen lain.
                        pass
                if low.endswith(".docx") and low.startswith("rk3"):
                    doc = Document(path)
                    for table in doc.tables:
                        for row in table.rows:
                            vals = [re.sub(r"\s+", " ", c.text or "").replace("\ufffd", " ").replace("\u2019", "'").strip() for c in row.cells]
                            if len(vals) >= 3 and re.match(r"^\d+", vals[0]) and vals[1] not in {"2", "Uraian Pekerjaan"}:
                                if vals[1] and vals[1] not in out["uraian_rk3"]:
                                    out["uraian_rk3"].append(vals[1])
                                if len(vals) >= 8 and "resiko paling tinggi" in vals[-1].lower():
                                    out["risiko_tertinggi"] = vals[2].lstrip("-• ").strip() or out["risiko_tertinggi"]
    except Exception:
        pass
    for name in os.listdir(folder):
        if name.lower().startswith("_hps_") and name.lower().endswith(".md"):
            try:
                for line in open(os.path.join(folder, name), encoding="utf-8"):
                    m = re.match(r"^\d+\s*\|\s*\*\*(.+?)\*\*\s*\|\s*-", line.strip())
                    if m and not m.group(1).startswith(("Jumlah", "Total", "Status", "Prompt")):
                        value = m.group(1).strip()
                        if value not in out["uraian"]:
                            out["uraian"].append(value)
            except Exception:
                pass
    if out["uraian_rk3"]:
        out["risiko"] = ", ".join(out["uraian_rk3"][:6])
    return out


def _write_local_enrichment(ws, enrichment: dict) -> None:
    """Tulis enrichment ke posisi Master Data PLPK yang stabil."""
    person = enrichment.get("personil") or []
    # Jangan sisakan alat/risiko dari template atau percobaan sebelumnya.
    for row in range(39, 57):
        ws.Cells(row, 3).ClearContents()
    ws.Cells(63, 3).ClearContents()
    ws.Cells(64, 3).ClearContents()
    if enrichment.get("contract_type"):
        ws.Cells(18, 3).Value = enrichment["contract_type"]
    if enrichment.get("cara_pembayaran"):
        ws.Cells(28, 3).Value = enrichment["cara_pembayaran"]
    for i, item in enumerate(person[:2]):
        base = 33 + i * 3
        ws.Cells(base, 3).Value = item.get("jabatan", "")
        ws.Cells(base + 1, 3).Value = item.get("pengalaman", "") or "0 Tahun"
        cert = item.get("sertifikat", "")
        jab = item.get("jabatan", "")
        if "k3" in jab.lower():
            cert = _normalize_k3_certificate(cert, jab)
        elif cert and "sk" not in cert.lower() and "pelaksana" in jab.lower():
            cert = "SKK " + cert
        ws.Cells(base + 2, 3).Value = cert
    for i, item in enumerate((enrichment.get("alat") or [])[:6]):
        ws.Cells(39 + i, 3).Value = item.get("nama", "")
        ws.Cells(45 + i, 3).Value = item.get("kapasitas", "")
        ws.Cells(51 + i, 3).Value = normalize_equipment_quantity(item.get("jumlah", ""))
    for row in range(66, 76):
        ws.Cells(row, 3).ClearContents()
    for i, value in enumerate((enrichment.get("uraian") or [])[:10]):
        ws.Cells(66 + i, 3).Value = value
    if enrichment.get("risiko"):
        ws.Cells(63, 3).Value = enrichment["risiko"]
    if enrichment.get("risiko_tertinggi"):
        ws.Cells(64, 3).Value = enrichment["risiko_tertinggi"]
    provider = enrichment.get("provider") or {}
    if provider.get("nama_penyedia"):
        ws.Cells(77, 3).Value = provider["nama_penyedia"]
    if provider.get("npwp_penyedia"):
        ws.Cells(78, 3).NumberFormat = "@"
        ws.Cells(78, 3).Value = provider["npwp_penyedia"]
    if provider.get("nomor_nota_dinas"):
        ws.Cells(87, 3).NumberFormat = "@"
        ws.Cells(87, 3).Value = provider["nomor_nota_dinas"]
    if provider.get("tgl_nota_dinas"):
        ws.Cells(88, 3).NumberFormat = "@"
        raw_date = str(provider["tgl_nota_dinas"])
        try:
            from datetime import date
            y, m, d = [int(x) for x in raw_date[:10].split("-")]
            months = ("", "Januari", "Februari", "Maret", "April", "Mei", "Juni", "Juli", "Agustus", "September", "Oktober", "November", "Desember")
            raw_date = f"{d} {months[m]} {y}"
        except Exception:
            pass
        ws.Cells(88, 3).Value = raw_date


def _apply_post_macro_enrichment(wb, excel_path: str, log) -> None:
    """Terapkan data lokal setelah macro inti selesai, untuk semua flow PLPK."""
    enrichment = _parse_local_enrichment(os.path.dirname(os.path.abspath(excel_path)))
    _write_local_enrichment(wb.Sheets("@ Master Data"), enrichment)
    log("Data lokal: personel, alat, uraian, risiko, pembayaran, ND disinkronkan.")


def _find_master_data_v2_root() -> str:
    """Cari clone ``procurement_core`` yang berisi modul snapshot V2.

    ``POKJA_CODE_ROOT`` pada sebagian setup menunjuk langsung ke clone
    ``Asisten_Pokja``. Modul V2 berada di clone sibling ``procurement_core``,
    sehingga resolver wajib mempertimbangkan ``POKJA_V19_ROOT`` dan parent
    dari root aplikasi.
    """
    app_root = os.path.dirname(os.path.abspath(__file__))
    code_root = str(os.environ.get("POKJA_CODE_ROOT") or "").strip()
    candidates = [
        os.environ.get("POKJA_V19_ROOT"),
        os.path.join(code_root, "procurement_core") if code_root else "",
        os.path.join(os.path.dirname(code_root), "procurement_core") if code_root else "",
        os.path.join(os.path.dirname(app_root), "procurement_core"),
    ]
    seen = set()
    for candidate in candidates:
        if not candidate:
            continue
        root = os.path.abspath(os.path.normpath(candidate))
        if root in seen:
            continue
        seen.add(root)
        if os.path.isfile(os.path.join(root, "master_data_v2.py")):
            return root
    return ""


def _sync_master_data_v2(excel_path: str, log) -> None:
    """Sinkronkan snapshot V2, dengan error tetap non-blocking."""
    core_root = _find_master_data_v2_root()
    if not core_root:
        log("[WARN] Snapshot V2 tidak tersinkron: master_data_v2.py tidak ditemukan")
        return

    import importlib
    import sys

    if core_root not in sys.path:
        sys.path.insert(0, core_root)
    expected_module = os.path.normcase(
        os.path.abspath(os.path.join(core_root, "master_data_v2.py"))
    )
    loaded = sys.modules.get("master_data_v2")
    loaded_file = (
        os.path.normcase(os.path.abspath(str(getattr(loaded, "__file__", "") or "")))
        if loaded
        else ""
    )
    if loaded and loaded_file != expected_module:
        sys.modules.pop("master_data_v2", None)
    module = importlib.import_module("master_data_v2")
    module.sync_daftar_paket_snapshot(excel_path)
    log("Snapshot Daftar Paket V2 tersinkron.")


def tulis_identitas_penyedia_ke_excel(
    excel_path: str,
    nama_penyedia: str,
    npwp_penyedia: str,
    progress_cb=None,
    *,
    name_cell: str = "C51",
    npwp_cell: str = "C52",
) -> dict:
    """Tulis identitas penyedia ke cell authoritative ``@ Master Data`` via COM.

    Default C51:C52 menjaga kompatibilitas PLJKK. PLPK memanggil helper ini
    dengan C77:C78.
    """
    def _log(message):
        if progress_cb:
            progress_cb(message)

    excel_path = os.path.abspath(excel_path)
    if not os.path.isfile(excel_path):
        return {"ok": False, "pesan": f"File tidak ditemukan: {excel_path}"}
    if not str(nama_penyedia or "").strip() and not str(npwp_penyedia or "").strip():
        return {"ok": False, "pesan": "Nama dan NPWP penyedia kosong."}

    import pythoncom
    import pywintypes
    import win32com.client

    pythoncom.CoInitialize()
    xl = None
    wb = None
    try:
        xl = win32com.client.DispatchEx("Excel.Application")
        xl.Visible = False
        xl.DisplayAlerts = False
        try:
            xl.AutomationSecurity = 1
        except Exception as exc:
            raise RuntimeError(
                "Excel AutomationSecurity gagal diatur ke macro-enabled; "
                "writer dibatalkan agar cache UDF tidak rusak."
            ) from exc
        xl.EnableEvents = False
        _log(f"Membuka Excel: {os.path.basename(excel_path)}")
        wb = xl.Workbooks.Open(
            excel_path,
            UpdateLinks=0,
            ReadOnly=False,
            IgnoreReadOnlyRecommended=True,
            AddToMru=False,
        )
        if bool(wb.ReadOnly):
            return {"ok": False, "pesan": "Workbook terbuka ReadOnly; tutup Excel paket lalu ulangi."}
        ws = wb.Worksheets("@ Master Data")
        ws.Range(name_cell).Value = str(nama_penyedia or "").strip()
        ws.Range(npwp_cell).NumberFormat = "@"
        ws.Range(npwp_cell).Value = str(npwp_penyedia or "").strip()
        wb.Save()

        saved_name = str(ws.Range(name_cell).Value or "").strip()
        saved_npwp = str(ws.Range(npwp_cell).Value or "").strip()
        if saved_name != str(nama_penyedia or "").strip() or saved_npwp != str(npwp_penyedia or "").strip():
            return {"ok": False, "pesan": f"Verifikasi {name_cell}:{npwp_cell} setelah Save tidak cocok."}
        _log(f"@ Master Data {name_cell}:{npwp_cell} tersimpan dan terverifikasi.")
        return {"ok": True, "pesan": f"Nama/NPWP penyedia tersimpan ke {name_cell}:{npwp_cell}."}
    except pywintypes.com_error as exc:
        return {"ok": False, "pesan": f"Excel COM error: {exc}"}
    except Exception as exc:
        return {"ok": False, "pesan": str(exc)}
    finally:
        if wb is not None:
            try:
                wb.Close(SaveChanges=False)
            except Exception:
                pass
        if xl is not None:
            try:
                xl.Quit()
            except Exception:
                pass
        try:
            pythoncom.CoUninitialize()
        except Exception:
            pass


def isi_master_data_pl(kode_paket: str, excel_path: str, progress_cb=None) -> dict:
    """Buka xlsm via COM, jalankan macro IsiDataPLByKode(kode_paket) dalam silent mode.

    Return: {"ok": bool, "pesan": str}
    """
    def _log(m):
        if progress_cb:
            progress_cb(m)

    excel_path = os.path.abspath(excel_path)
    if not os.path.isfile(excel_path):
        return {"ok": False, "pesan": f"File tidak ditemukan: {excel_path}"}
    if not kode_paket:
        return {"ok": False, "pesan": "kode_paket kosong"}

    import win32com.client
    import pythoncom
    import pywintypes
    pythoncom.CoInitialize()

    xl = None
    wb = None
    sync_v2 = False
    try:
        xl = win32com.client.DispatchEx("Excel.Application")
        xl.Visible = False
        xl.DisplayAlerts = False
        # AutomationSecurity=1 (msoAutomationSecurityLow) — izinkan macro jalan
        try:
            xl.AutomationSecurity = 1
        except Exception as exc:
            raise RuntimeError(
                "Excel AutomationSecurity gagal diatur ke macro-enabled; "
                "writer dibatalkan agar cache UDF tidak rusak."
            ) from exc
        xl.EnableEvents = False

        _log(f"Membuka Excel: {os.path.basename(excel_path)}")
        wb = xl.Workbooks.Open(excel_path, UpdateLinks=0)

        # Aktifkan silent mode → suppress MsgBox di jalur VBA
        try:
            xl.Run("ModDraftPaketPL.SetSilentPL", True)
        except pywintypes.com_error as ce:
            return {"ok": False, "pesan": f"Macro SetSilentPL tidak ditemukan/compile error: {ce}"}

        _log(f"Mengisi @ Master Data untuk {kode_paket}...")
        try:
            xl.Run("ModDraftPaketPL.IsiDataPLByKode", str(kode_paket))
        except pywintypes.com_error as ce:
            return {"ok": False, "pesan": f"Macro IsiDataPLByKode gagal: {ce}"}

        # Jalur tombol "Isi Excel" harus identik dengan bulk-create untuk
        # PLPK. Jangan menerapkan mapping PLPK pada workbook PLJKK.
        if _is_plpk_workbook(excel_path):
            try:
                _apply_post_macro_enrichment(wb, excel_path, _log)
            except Exception as local_e:
                _log(f"[WARN] Data lokal: {local_e}")

        # Refresh sheet @ Evaluasi (tgl_pembukaan, nomor BA, dll) — 1 sesi COM
        try:
            xl.Run("ModDraftPaketPL.IsiEvaluasiPLStandalone")
            _log("@ Evaluasi ter-refresh.")
        except pywintypes.com_error:
            _log("[WARN] IsiEvaluasiPLStandalone tidak ditemukan — skip.")

        wb.Save()
        _log("@ Master Data + @ Evaluasi terisi.")
        sync_v2 = True
        return {"ok": True, "pesan": "@ Master Data + @ Evaluasi terisi otomatis"}
    except pywintypes.com_error as ce:
        return {"ok": False, "pesan": f"COM error: {ce}"}
    except Exception as e:
        return {"ok": False, "pesan": str(e)}
    finally:
        if wb is not None:
            try:
                wb.Close(SaveChanges=False)
            except Exception:
                pass
        if xl is not None:
            try:
                xl.Quit()
            except Exception:
                pass
        if sync_v2:
            try:
                _sync_master_data_v2(excel_path, _log)
            except Exception as exc:
                _log(f"[WARN] Snapshot V2 tidak tersinkron: {exc}")
        try:
            pythoncom.CoUninitialize()
        except Exception:
            pass


# ── CLI self-test ─────────────────────────────────────────────

def proses_hps_dan_master_data(kode_paket: str, excel_path: str,
                                hps_hasil: dict = None,
                                progress_cb=None,
                                timeout: int = 90,
                                jenis_pl: str = "",
                                tanggal_create=None) -> dict:
    """1 sesi COM: tulis HPS (jika ada) lalu IsiDataPLByKode — 1x DispatchEx.

    Dipakai oleh _proses_excel_paket_pl() di app.py (bulk-create folder PL).
    Urutan: Open -> (a) tulis sheet '5. HPS' [isolasi, gagal tidak blokir] ->
    (b) SetSilentPL(True) -> IsiDataPLByKode(kode) -> Save -> finally Close+Quit.

    Args:
        kode_paket: kode paket PL (string).
        excel_path: path absolut file .xlsm yang sudah di-refresh.
        hps_hasil:  dict {items, total_nilai, total_nilai_bulat} dari scrape_hps_pl().
                    Jika None, langkah HPS dilewati.
        progress_cb: callable(str) opsional untuk log progres.

    Return: {"ok": bool, "hps": {"ok", "pesan", "count"}, "md": {"ok", "pesan"}}
    """
    def _log(m):
        if progress_cb:
            try:
                progress_cb(m)
            except Exception:
                pass

    excel_path = os.path.abspath(excel_path)
    if not os.path.isfile(excel_path):
        return {
            "ok": False,
            "hps": {"ok": False, "pesan": "file tidak ada", "count": 0},
            "md":  {"ok": False, "pesan": f"File tidak ditemukan: {excel_path}"},
        }
    if not kode_paket:
        return {
            "ok": False,
            "hps": {"ok": False, "pesan": "dilewati", "count": 0},
            "md":  {"ok": False, "pesan": "kode_paket kosong"},
        }

    result_box = [None]
    import threading

    def _run_com():
        import win32com.client
        import pythoncom
        import pywintypes
        import hps_engine as _hps_eng

        pythoncom.CoInitialize()

        xl = None
        wb = None
        hps_res = {"ok": False, "pesan": "dilewati", "count": 0}
        md_res  = {"ok": False, "pesan": ""}

        try:
            xl = win32com.client.DispatchEx("Excel.Application")
            xl.Visible = False
            xl.DisplayAlerts = False
            try:
                xl.AutomationSecurity = 1  # msoAutomationSecurityLow
            except Exception as exc:
                raise RuntimeError(
                    "Excel AutomationSecurity gagal diatur ke macro-enabled; "
                    "writer dibatalkan agar cache UDF tidak rusak."
                ) from exc
            xl.EnableEvents = False

            _log(f"Membuka Excel: {os.path.basename(excel_path)}")
            wb = xl.Workbooks.Open(excel_path, UpdateLinks=0)

            # (a) Tulis HPS ke sheet '5. HPS' --- ISOLASI: gagal tidak blokir Master Data
            if hps_hasil and hps_hasil.get("items"):
                try:
                    ws_hps = wb.Sheets("5. HPS")
                    r = _hps_eng._tulis_hps_ke_ws(ws_hps, wb, hps_hasil, progress_cb)
                    hps_res = {"ok": r["ok"], "pesan": r["pesan"], "count": r.get("count", 0)}
                    _log(f"HPS: {hps_res['count']} baris ditulis.")
                except Exception as e_hps:
                    hps_res = {"ok": False, "pesan": f"Gagal tulis HPS: {e_hps}", "count": 0}
                    _log(f"WARN HPS: {e_hps} (Master Data tetap lanjut)")
            else:
                hps_res = {"ok": True, "pesan": "hps_hasil kosong, dilewati", "count": 0}

            # (b) SetSilentPL + IsiDataPLByKode
            try:
                xl.Run("ModDraftPaketPL.SetSilentPL", True)
            except pywintypes.com_error as ce:
                md_res = {"ok": False, "pesan": f"Macro SetSilentPL tidak ditemukan/compile error: {ce}"}
                wb.Save()
                result_box[0] = {"ok": False, "hps": hps_res, "md": md_res}
                return

            _log(f"Mengisi @ Master Data untuk {kode_paket}...")
            try:
                xl.Run("ModDraftPaketPL.IsiDataPLByKode", str(kode_paket))
                md_res = {"ok": True, "pesan": "@ Master Data terisi otomatis"}
            except pywintypes.com_error as ce:
                md_res = {"ok": False, "pesan": f"Macro IsiDataPLByKode gagal: {ce}"}

            # Enrichment lokal wajib dilakukan setelah macro: Supabase hanya
            # membawa field inti, sedangkan alat/uraian/RK3 berada di paket.
            if md_res["ok"] and _is_plpk_workbook(excel_path, jenis_pl):
                try:
                    _apply_post_macro_enrichment(wb, excel_path, _log)
                except Exception as local_e:
                    _log(f"WARN data lokal: {local_e}")
                # Saat create-folder bulk/single, HPS sudah tersedia di sesi
                # COM yang sama. Tulis uraian final langsung ke C19 agar tidak
                # bergantung pada cache/upsert Supabase yang mungkin belum ada.
                if hps_hasil and hps_hasil.get("items"):
                    try:
                        uraian = _hps_eng._build_uraian_singkat_pk(
                            hps_hasil["items"], excel_path
                        )
                        if uraian:
                            ws_master = wb.Sheets("@ Master Data")
                            ws_master.Cells(19, 3).NumberFormat = "General"
                            ws_master.Cells(19, 3).Value = uraian
                            _log("Uraian singkat C19 dibuat dari DIVISI HPS.")
                    except Exception as uraian_e:
                        _log(f"WARN uraian singkat C19: {uraian_e}")

            # Macro membaca cache Supabase lebih dulu. Normalisasi di boundary
            # workbook memastikan boilerplate tender "gabungan" tidak masuk
            # ke Excel PLPK, termasuk saat sumber lokal tidak memuat field ini.
            if md_res["ok"]:
                try:
                    ws_master = wb.Sheets("@ Master Data")
                    current_contract = ws_master.Cells(18, 3).Value
                    normalized_contract = _normalize_contract_type(current_contract)
                    if normalized_contract and normalized_contract != current_contract:
                        ws_master.Cells(18, 3).Value = normalized_contract
                        _log("Jenis kontrak PLPK dinormalisasi: Harga Satuan.")
                except Exception as contract_e:
                    _log(f"WARN normalisasi jenis kontrak: {contract_e}")

            # Create-folder menetapkan tanggal BA Reviu dan Dokpil sekali,
            # setelah macro selesai menyalin cache SPSE yang bisa berupa teks.
            # Jika boundary ini gagal, jangan simpan workbook parsial.
            if md_res["ok"] and tanggal_create is not None:
                try:
                    set_create_folder_dates(wb, tanggal_create, _log)
                except Exception as date_e:
                    _log(f"WARN tanggal create-folder: {date_e}")
                    raise RuntimeError(
                        f"tanggal create-folder gagal diverifikasi: {date_e}"
                    ) from date_e

            # Refresh @ Evaluasi setelah Master Data terisi
            if md_res["ok"]:
                try:
                    xl.Run("ModDraftPaketPL.IsiEvaluasiPLStandalone")
                    _log("@ Evaluasi ter-refresh.")
                except pywintypes.com_error:
                    _log("[WARN] IsiEvaluasiPLStandalone tidak ditemukan — skip.")

            wb.Save()
            _log("Excel disimpan.")
            result_box[0] = {"ok": md_res["ok"], "hps": hps_res, "md": md_res}

        except Exception as e:
            result_box[0] = {
                "ok": False,
                "hps": hps_res,
                "md":  {"ok": False, "pesan": str(e)},
            }
        finally:
            if wb is not None:
                try:
                    wb.Close(SaveChanges=False)
                except Exception:
                    pass
            if xl is not None:
                try:
                    xl.Quit()
                except Exception:
                    pass
            try:
                pythoncom.CoUninitialize()
            except Exception:
                pass

    t = threading.Thread(target=_run_com, daemon=True)
    t.start()
    t.join(timeout=timeout)
    if t.is_alive():
        return {
            "ok": False,
            "hps": {"ok": False, "pesan": "timeout", "count": 0},
            "md":  {"ok": False, "pesan": f"COM timeout {timeout}s — Excel tidak responsif"},
        }
    result = result_box[0] or {
        "ok": False,
        "hps": {"ok": False, "pesan": "tidak ada hasil", "count": 0},
        "md":  {"ok": False, "pesan": "thread selesai tanpa hasil"},
    }
    if result.get("md", {}).get("ok"):
        try:
            _sync_master_data_v2(excel_path, _log)
        except Exception as exc:
            _log(f"[WARN] Snapshot V2 tidak tersinkron: {exc}")
    return result

if __name__ == "__main__":
    import sys
    if len(sys.argv) < 3:
        print("Usage: python isi_master_data_pl.py <kode_paket> <path.xlsm>")
        sys.exit(1)
    res = isi_master_data_pl(sys.argv[1], sys.argv[2], progress_cb=print)
    print(res)
