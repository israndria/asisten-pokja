"""Rebuild template undangan PL dari nol."""
import docx
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = (
    r"D:\Dokumen\@ POKJA 2026"
    r"\Paket Experiment - Pengadaan Langsung"
    r"\Development - PL - PK\4. Undangan Full PLPK - Template.docx"
)

doc = Document()

# Margin
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.0)

# Default style
style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(12)


def set_font(run, size=12, bold=False):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold = bold


# Judul
p_judul = doc.add_paragraph()
p_judul.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_judul.paragraph_format.space_before = Pt(0)
p_judul.paragraph_format.space_after = Pt(6)
run_j = p_judul.add_run("SURAT UNDANGAN REVIU")
set_font(run_j, size=12, bold=True)

# Tabel 4 kolom: label | sep | isi | kanan
t = doc.add_table(rows=0, cols=4)
tbl = t._tbl

# Hapus semua border
tblPr = tbl.find(qn("w:tblPr"))
if tblPr is None:
    tblPr = OxmlElement("w:tblPr")
    tbl.insert(0, tblPr)
tblBorders = OxmlElement("w:tblBorders")
for side in ["top", "left", "bottom", "right", "insideH", "insideV"]:
    b = OxmlElement(f"w:{side}")
    b.set(qn("w:val"), "none")
    tblBorders.append(b)
tblPr.append(tblBorders)

W4 = [3.0, 0.5, 7.5, 3.5]  # cm: label | sep | isi | kanan


def r4(
    c0="", c1="", c2="", c3="",
    bold=False, size=12,
    a0=WD_ALIGN_PARAGRAPH.LEFT,
    a1=WD_ALIGN_PARAGRAPH.LEFT,
    a2=WD_ALIGN_PARAGRAPH.LEFT,
    a3=WD_ALIGN_PARAGRAPH.LEFT,
):
    row = t.add_row()
    for i, (cell, txt, al) in enumerate(zip(row.cells, [c0, c1, c2, c3], [a0, a1, a2, a3])):
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        p = cell.paragraphs[0]
        p.alignment = al
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(str(txt))
        set_font(run, size=size, bold=bold)
        # Set lebar kolom
        tc = cell._tc
        tcPr = tc.find(qn("w:tcPr"))
        if tcPr is None:
            tcPr = OxmlElement("w:tcPr")
            tc.insert(0, tcPr)
        tcW = OxmlElement("w:tcW")
        tcW.set(qn("w:w"), str(int(W4[i] * 567)))
        tcW.set(qn("w:type"), "dxa")
        ex = tcPr.find(qn("w:tcW"))
        if ex is not None:
            tcPr.remove(ex)
        tcPr.append(tcW)
    return row


# ── Isi surat ──────────────────────────────────────────────────────────────────

# Tanggal di kanan
r4(c3="Rantau, «TANGGAL_KIRIM»", a3=WD_ALIGN_PARAGRAPH.RIGHT)
r4()
# Header surat
r4("Nomor",    ":", "«NOMOR_SURAT»")
r4("Lampiran", ":", "1 (satu) berkas")
r4("Perihal",  ":", "Undangan Reviu Dokumen Persiapan Pengadaan")
r4()
# Kepada Yth di kanan
r4(c3="Kepada Yth.")
r4(c3="«PENERIMA_1»")
r4(c3="Di -")
r4(c3="Tempat")
r4()
# Badan surat
r4("Dengan hormat,")
r4()
r4("Berdasarkan:")
r4(
    "1.", "",
    "Peraturan Lembaga Kebijakan Pengadaan Barang/Jasa Pemerintah Nomor 12 "
    "Tahun 2021 Tentang Pedoman Pelaksanaan Pengadaan Barang/Jasa Melalui "
    "Penyedia, Lampiran I BAB III Angka 3.1 tentang Reviu Dokumen Persiapan "
    "Pengadaan;",
)
r4("2.", "", "«BERDASARKAN_PP»")
r4()
r4("Diharapkan kehadirannya pada:")
r4("Hari / Tanggal", ":", "«HARI_TGL_RAPAT»")
r4("Pukul",          ":", "«PUKUL»")
r4("Acara",          ":", "«ACARA»")
r4("Dokumen Yang Dibawa", ":", "1.   Spesifikasi Teknis / KAK")
r4("", "", "2.   Dokumen HPS")
r4("", "", "3.   Rancangan Kontrak")
r4("", "", "4.   Dokumen Anggaran Belanja")
r4("Tempat", ":", "«TEMPAT_RAPAT»")
r4()
r4("Demikian kami sampaikan, atas kehadirannya kami ucapkan terima kasih.")
r4()
# Blok TTD
r4("«INSTANSI_PP»")
r4()
r4("«TTD_PP»")   # ← engine ganti dengan gambar
r4()
r4()
r4("«NAMA_NIP_PP»")

doc.save(OUT)
print("Tersimpan:", OUT)

# Verifikasi
doc2 = Document(OUT)
t2 = doc2.tables[0]
seen = set()
for ri, row in enumerate(t2.rows):
    uniq = []
    for cell in row.cells:
        cid = id(cell._tc)
        if cid in seen:
            continue
        seen.add(cid)
        txt = cell.text.strip()
        if txt:
            uniq.append(txt[:70])
    if uniq:
        print(f"R{ri}: {uniq}")
